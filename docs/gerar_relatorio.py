#!/usr/bin/env python3
"""
Gera docs/relatorio_astar_abnt.docx — ABNT NBR 14724.

Melhorias v2:
  • Margens ABNT exatas: topo/esq 3 cm, baixo/dir 2 cm
  • Espaçamento 1,5 (ABNT)
  • Recuo de primeira linha 1,25 cm (ABNT) — não recuo de bloco
  • Títulos e subtítulos numerados explicitamente (sem numPr automático)

Uso (raiz do repositório):
    python docs/gerar_relatorio.py
"""
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT     = Path(__file__).resolve().parent.parent
GRAFICOS = ROOT / "results" / "graficos"
MODELO   = ROOT / "docs" / "relatorio_dijkstra_abnt.docx"
SAIDA    = ROOT / "docs" / "relatorio_astar_abnt.docx"

# ── Template ─────────────────────────────────────────────────────────────────
shutil.copy(MODELO, SAIDA)
doc = Document(str(SAIDA))

body = doc.element.body
for child in list(body)[:-1]:       # preserva o sectPr
    body.remove(child)

# ── Margens ABNT (NBR 14724): topo e esq 3 cm; baixo e dir 2 cm ──────────────
sec = doc.sections[0]
sec.top_margin    = Cm(3)
sec.bottom_margin = Cm(2)
sec.left_margin   = Cm(3)
sec.right_margin  = Cm(2)

# ── Constantes de formatação ──────────────────────────────────────────────────
LS_15  = 360   # 1,5 × 240 = 360 (valor w:line para espaçamento 1,5)
LS_10  = 240   # espaçamento simples
FIRST  = 709   # 1,25 cm em twips  (1,25 × 1440/2,54 ≈ 709)

# ── Helpers XML ───────────────────────────────────────────────────────────────

def _spacing(p, line=LS_15, rule="auto", before=0, after=0):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    s = OxmlElement("w:spacing")
    s.set(qn("w:line"),     str(line))
    s.set(qn("w:lineRule"), rule)
    if before: s.set(qn("w:before"), str(before))
    if after:  s.set(qn("w:after"),  str(after))
    pPr.append(s)

def _indent(p, first_line=0, left=0, right=0, hanging=0):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    if not any([first_line, left, right, hanging]):
        return
    ind = OxmlElement("w:ind")
    if first_line: ind.set(qn("w:firstLine"), str(first_line))
    if left:       ind.set(qn("w:left"),       str(left))
    if right:      ind.set(qn("w:right"),      str(right))
    if hanging:    ind.set(qn("w:hanging"),    str(hanging))
    pPr.append(ind)

def _jc(p, val="both"):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:jc")):
        pPr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), val)
    pPr.append(jc)

def _disable_numbering(p):
    """Remove numeração automática herdada do estilo."""
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:numPr")):
        pPr.remove(old)
    numPr = OxmlElement("w:numPr")
    ilvl  = OxmlElement("w:ilvl");  ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "0")
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)

def _run(p, text, size=12, bold=False, italic=False, font="Times New Roman"):
    run = p.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    return run

def _table_borders(tbl):
    tbl_el = tbl._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)

def _cell_shade(cell, fill="E1E1E1"):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)

# ── Parágrafos de conteúdo ────────────────────────────────────────────────────

def body_text(text, bold=False, italic=False):
    p = doc.add_paragraph(style="Body Text")
    _spacing(p, line=LS_15, before=0, after=0)
    _indent(p, first_line=FIRST)
    _jc(p, "both")
    _run(p, text, bold=bold, italic=italic)
    return p

def body_mixed(parts):
    """parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph(style="Body Text")
    _spacing(p, line=LS_15, before=0, after=0)
    _indent(p, first_line=FIRST)
    _jc(p, "both")
    for text, bold, italic in parts:
        _run(p, text, bold=bold, italic=italic)
    return p

def h1(num, text):
    """Título de 1.º nível: '1 INTRODUÇÃO'"""
    p = doc.add_paragraph(style="Heading1")
    _disable_numbering(p)
    _spacing(p, line=LS_15, before=240, after=0)
    _indent(p)
    _jc(p, "left")
    _run(p, f"{num} {text.upper()}")
    return p

def h2(num, text):
    """Título de 2.º nível: '2.1 Definição e objetivo'"""
    p = doc.add_paragraph(style="Heading2")
    _disable_numbering(p)
    _spacing(p, line=LS_15, before=120, after=0)
    _indent(p)
    _jc(p, "left")
    run = p.add_run(f"{num} {text}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def list_item(text):
    p = doc.add_paragraph(style="List Paragraph")
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:numPr")):
        pPr.remove(old)
    _spacing(p, line=LS_15, before=0, after=0)
    _indent(p, left=FIRST, hanging=FIRST)   # recuo deslocado para bullet
    _jc(p, "both")
    _run(p, f"– {text}")
    return p

def caption_table(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, line=LS_10, before=160, after=40)
    _indent(p)
    _run(p, text, size=10, bold=True)
    return p

def caption_figure(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, line=LS_10, before=40, after=0)
    _indent(p)
    _run(p, text, size=10, bold=True)
    return p

def source_line():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, line=LS_10, before=0, after=160)
    _indent(p)
    _run(p, "Fonte: elaborado pelos autores (2026).", size=10)
    return p

def figure(img_name, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, line=LS_10, before=160, after=0)
    _indent(p)
    p.add_run().add_picture(str(GRAFICOS / img_name), width=Cm(15))
    caption_figure(caption_text)
    source_line()

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _spacing(p, line=LS_10, before=0, after=0)
        _indent(p, left=FIRST)
        _run(p, line, size=10, font="Courier New")

def normal_center(text, size=12, bold=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, line=LS_15, before=before, after=after)
    _indent(p)
    if text:
        _run(p, text, size=size, bold=bold)
    return p

def page_break():
    p = doc.add_paragraph()
    _spacing(p, line=LS_10)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    p.add_run()._r.append(br)

def toc_entry(text, level=1):
    style = "TOC 1" if level == 1 else "TOC 2"
    p = doc.add_paragraph(style=style)
    _run(p, text)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Normal"
    _table_borders(t)
    # Cabeçalho
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        _cell_shade(cell, "E1E1E1")
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Paragraph"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, line=LS_10, before=60)
        _indent(p, left=11, right=1)
        _run(p, h, size=9, bold=True)
    # Dados
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Paragraph"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, line=LS_10, before=60)
            _indent(p, left=11, right=1)
            _run(p, str(val), size=9)
    return t

# ════════════════════════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
_spacing(p, line=LS_15)
run = p.add_run("CESAR SCHOOL\nCIÊNCIA DA COMPUTAÇÃO\nTEORIA DA COMPUTAÇÃO")
run.font.name = "Times New Roman"; run.font.size = Pt(12); run.bold = True

for _ in range(8): normal_center("")

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
_spacing(p_title, line=LS_15)
run = p_title.add_run("ALGORITMO A* (A-ESTRELA)")
run.font.name = "Times New Roman"; run.font.size = Pt(18); run.bold = True

normal_center("")

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
_spacing(p_sub, line=LS_15)
run = p_sub.add_run(
    "Análise teórica e experimental do caminho mínimo em grades:\n"
    "implementações em Python e JavaScript"
)
run.font.name = "Times New Roman"; run.font.size = Pt(14)

for _ in range(8): normal_center("")

p_auth = doc.add_paragraph()
p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
_spacing(p_auth, line=LS_15)
run = p_auth.add_run("Arthur Padilha   Pedro Tojal   Arthur Lins   Luis Guilherme")
run.font.name = "Times New Roman"; run.font.size = Pt(12)

normal_center("")
normal_center("Prof. Daniel Bezerra", size=12)

for _ in range(6): normal_center("")

normal_center("Recife", size=14, bold=True)
normal_center("2026",   size=14, bold=True)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SUMÁRIO
# ════════════════════════════════════════════════════════════════════════════

normal_center("SUMÁRIO", size=14, bold=True, after=280)

TOC = [
    (1, "1  Introdução .....................................................................................................  3"),
    (1, "2  O Algoritmo A* ................................................................................................  4"),
    (2, "     2.1  Definição e objetivo ....................................................................................  4"),
    (2, "     2.2  Funcionamento passo a passo ..........................................................................  4"),
    (2, "     2.3  Estrutura de dados: heap binário de mínimo ........................................................  5"),
    (2, "     2.4  Pseudocódigo ............................................................................................  6"),
    (2, "     2.5  Complexidade computacional (O, Ω e Θ) ..............................................................  6"),
    (1, "3  Metodologia Experimental ......................................................................................  8"),
    (2, "     3.1  Implementações ...........................................................................................  8"),
    (2, "     3.2  Cenários de teste e geração das entradas ............................................................  8"),
    (2, "     3.3  Tamanhos avaliados .......................................................................................  9"),
    (2, "     3.4  Protocolo de medição .....................................................................................  9"),
    (2, "     3.5  Ambiente de execução ....................................................................................  10"),
    (2, "     3.6  Resumo dos cenários executados ........................................................................  10"),
    (1, "4  Resultados e Análise dos Gráficos ...........................................................................  13"),
    (2, "     4.1  Curvas de tempo por cenário .............................................................................  13"),
    (2, "     4.2  Comparação entre Python e JavaScript ................................................................  16"),
    (2, "     4.3  Trabalho algorítmico: nós expandidos .................................................................  17"),
    (2, "     4.4  Validação do modelo teórico ............................................................................  18"),
    (1, "5  Aplicabilidade: Eficiência e Limitações ....................................................................  19"),
    (2, "     5.1  Contextos de eficiência e aplicações .................................................................  19"),
    (2, "     5.2  Limitações .................................................................................................  20"),
    (1, "6  Reflexão Teórica: Classes de Complexidade .................................................................  21"),
    (2, "     6.1  O problema pertence à classe P? .......................................................................  21"),
    (2, "     6.2  Existe uma versão NP do problema? ....................................................................  21"),
    (2, "     6.3  Problemas semelhantes NP-completos ...................................................................  22"),
    (1, "7  Execução do Projeto ............................................................................................  23"),
    (1, "8  Conclusão .........................................................................................................  24"),
    (1, "REFERÊNCIAS ..........................................................................................................  25"),
]
for level, text in TOC:
    toc_entry(text, level)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1 INTRODUÇÃO
# ════════════════════════════════════════════════════════════════════════════

h1("1", "Introdução")

body_text(
    "O problema do caminho mínimo é central na Teoria da Computação e em inúmeras aplicações "
    "práticas. Em contextos onde se dispõe de um destino conhecido e de uma estimativa do custo "
    "restante para alcançá-lo, a busca pode ser guiada de forma muito mais eficiente do que na "
    "exploração cega. O algoritmo A* (lê-se “A-estrela”), proposto por Peter Hart, Nils Nilsson "
    "e Bertram Raphael em 1968, formaliza essa intuição: ele ordena a fronteira de busca por uma "
    "função f(v) = g(v) + h(v), onde g(v) é o custo real acumulado da origem até v e h(v) é uma "
    "estimativa heurística do custo de v até o destino. Quando a heurística é admissível — nunca "
    "superestima o custo real — o A* é garantidamente ótimo: o primeiro caminho encontrado é o de "
    "menor custo. Quando é também consistente (monótona), cada vértice é expandido no máximo uma "
    "vez, eliminando retrabalho."
)
body_text(
    "Este trabalho tem como objetivo estudar o algoritmo A* sob duas perspectivas complementares. "
    "Na perspectiva teórica, descreve-se o que o algoritmo faz, como funciona passo a passo, qual "
    "a estrutura de dados empregada e qual a sua complexidade assintótica nas notações O, Ω e Θ. "
    "Na perspectiva experimental, o algoritmo foi implementado em duas linguagens de programação "
    "— Python e JavaScript (executado no Node.js) — e submetido a um conjunto de experimentos "
    "controlados com o objetivo de medir o tempo de execução em três cenários distintos e validar "
    "empiricamente o comportamento previsto pela teoria."
)
body_mixed([
    ("O domínio escolhido para os experimentos é uma ", False, False),
    ("grade 2D com obstáculos", True, False),
    (": cada célula livre é um vértice, vizinha das células livres adjacentes em quatro direções "
     "(cima, baixo, esquerda, direita), com custo unitário por passo. A heurística empregada é a ",
     False, False),
    ("distância de Manhattan", True, False),
    (", h(v) = |v.linha − destino.linha| + |v.coluna − destino.coluna|, que é admissível e "
     "consistente para movimento em quatro direções com custo unitário — garantindo a otimalidade "
     "do algoritmo.", False, False),
])
body_text(
    "O relatório está organizado da seguinte forma. A Seção 2 apresenta a fundamentação teórica "
    "do algoritmo. A Seção 3 detalha a metodologia experimental. A Seção 4 discute os resultados. "
    "A Seção 5 aborda a aplicabilidade. A Seção 6 traz a reflexão teórica sobre as classes de "
    "complexidade P e NP. A Seção 7 explica como executar os códigos e a Seção 8 conclui."
)

# ════════════════════════════════════════════════════════════════════════════
# 2 O ALGORITMO A*
# ════════════════════════════════════════════════════════════════════════════

h1("2", "O Algoritmo A*")

h2("2.1", "Definição e objetivo")

body_mixed([
    ("O algoritmo A* resolve o problema do ", False, False),
    ("caminho de menor custo entre dois vértices", True, False),
    (" de um grafo com pesos não negativos nas arestas, utilizando uma ", False, False),
    ("função heurística", True, False),
    (" que estima o custo restante até o destino. Enquanto o Dijkstra explora todos os vértices "
     "em ordem crescente de distância à origem — equivalente ao A* com h ≡ 0 —, o A* “puxa” a "
     "busca na direção do destino, expandindo primeiro os vértices com maior potencial de levar "
     "a um caminho ótimo. Isso pode reduzir drasticamente o número de vértices processados, "
     "especialmente quando a heurística é precisa.", False, False),
])
body_text("Duas propriedades da heurística são suficientes para garantir a otimalidade:")
list_item(
    "Admissibilidade: h(v) ≤ custo_real(v, destino) para todo v. A heurística nunca superestima, "
    "portanto o algoritmo nunca descarta prematuramente um caminho ótimo."
)
list_item(
    "Consistência (monotonicidade): h(u) ≤ custo(u,v) + h(v) para toda aresta (u,v). Com "
    "consistência, cada vértice é expandido no máximo uma vez."
)
body_text(
    "A distância de Manhattan satisfaz ambas as propriedades para movimento em quatro direções "
    "com custo unitário: cada passo reduz a distância de Manhattan em exatamente 1 (direção "
    "correta) ou aumenta em 1 (perpendicular), portanto a estimativa nunca excede o custo real."
)

h2("2.2", "Funcionamento passo a passo")

body_text("O funcionamento do A* pode ser resumido nos seguintes passos:")
list_item(
    "Inicialização: g(origem) ← 0; origem inserida na lista aberta com f = h(origem); "
    "demais g-scores considerados infinitos."
)
list_item(
    "Extração do mínimo: remove-se da lista aberta o vértice v com menor f. Se já foi "
    "expandido (entrada obsoleta — lazy deletion), descarta-se e continua."
)
list_item(
    "Verificação do objetivo: se v = destino, retorna o caminho reconstruído pelos pais e g(v)."
)
list_item(
    "Expansão dos vizinhos: para cada vizinho u livre de v, calcula-se g' = g(v) + 1. "
    "Se g' < g(u), atualiza-se g(u), registra-se v como pai de u e insere-se u na lista "
    "aberta com f = g' + h(u)."
)
list_item(
    "Repetição: extração e expansão se repetem até a lista ficar vazia (sem caminho) ou "
    "o destino ser extraído."
)
body_mixed([
    ("A ", False, False),
    ("lazy deletion", True, False),
    (" evita remoções custosas do meio do heap: entradas obsoletas são descartadas ao serem "
     "extraídas, preservando a complexidade assintótica e simplificando a implementação.",
     False, False),
])

h2("2.3", "Estrutura de dados: heap binário de mínimo")

body_mixed([
    ("A eficiência do A* depende da estrutura usada para extrair o vértice com menor f. "
     "Utilizou-se uma ", False, False),
    ("fila de prioridade implementada como heap binário de mínimo", True, False),
    (", com inserção e remoção em O(log n).", False, False),
])
body_text(
    "O critério de desempate possui três níveis: (1) menor f = g + h; (2) menor h (vértice "
    "mais próximo do destino); (3) menor ordem de inserção (determinismo total). O desempate "
    "determinístico é essencial para a comparabilidade entre linguagens."
)
body_mixed([
    ("Em ", False, False),
    ("Python", True, False),
    (" utilizou-se o módulo heapq. Em ", False, False),
    ("JavaScript", True, False),
    (" desenvolveu-se uma classe MinHeap própria (bubble up / sink down). Ambas as "
     "implementações usam Map/Set para g-scores, pais e conjunto de expandidos — O(1) médio.",
     False, False),
])

h2("2.4", "Pseudocódigo")

body_text("A seguir, o pseudocódigo do algoritmo implementado:")
code_block([
    "A*(grade, origem, destino):",
    "    aberta  ← heap-mínimo ordenado por (f, h, ordem de inserção)",
    "    g[origem] ← 0",
    "    insere origem na aberta com f = h(origem)",
    "",
    "    enquanto aberta não vazia:",
    "        (f, h, _, v) ← extrai_minimo(aberta)",
    "        se v já foi expandido: continua    # lazy deletion",
    "        marca v como expandido",
    "",
    "        se v = destino:",
    "            retorna (caminho reconstruído pelos pais, g[v])",
    "",
    "        para cada vizinho u de v (4 direções, livre, dentro da grade):",
    "            g' ← g[v] + 1",
    "            se g' < g[u]:",
    "                g[u] ← g'",
    "                pai[u] ← v",
    "                hu ← manhattan(u, destino)",
    "                insere u na aberta com f = g' + hu",
    "",
    "    retorna \"sem caminho\"",
])

h2("2.5", "Complexidade computacional (O, Ω e Θ)")

body_mixed([
    ("Define-se ", False, False),
    ("n", True, False),
    (" como o número total de células (n = lado²). Cada célula tem no máximo 4 vizinhos "
     "(E = O(n)); com heap binário, cada operação custa O(log n).", False, False),
])
body_mixed([
    ("Limite superior — O(n log n):", True, False),
    (" no pior caso todos os n vértices passam pelo heap (até 4 inserções por vértice). "
     "Custo total O(n log n). Nenhuma execução ultrapassa esse limite.", False, False),
])
body_mixed([
    ("Limite inferior — Ω(L·log n):", True, False),
    (" o algoritmo precisa expandir cada vértice do caminho ótimo de comprimento L. "
     "No melhor caso L ≈ √n, resultando em Ω(√n·log n), empiricamente Θ(√n).", False, False),
])
body_mixed([
    ("Limite justo — Θ por caso:", True, False),
    (" não existe Θ global, pois os limites superior e inferior não coincidem para todo "
     "conjunto de entradas. A análise por caso é:", False, False),
])
list_item(
    "Melhor caso (grade vazia, caminho reto): heurística perfeita; A* expande exatamente "
    "√n + 1 vértices. Tempo Θ(√n·log n), empiricamente Θ(√n). Confirmado: 51, 151 e 301 "
    "expansões para grades 51×51, 151×151 e 301×301."
)
list_item(
    "Caso médio (obstáculos aleatórios 25%): ~8% das células expandidas, fração "
    "aproximadamente constante em n. Tempo ≈ Θ(n·log n) com constante pequena, "
    "empiricamente Θ(n)."
)
list_item(
    "Pior caso (labirinto em serpentina): ~n/2 vértices expandidos. Tempo Θ(n log n). "
    "1.301, 11.401 e 45.601 expansões para n = 2.601, 22.801 e 90.601."
)
body_mixed([
    ("Complexidade de espaço:", True, False),
    (" Θ(n) em todos os casos — g-scores, pais e heap crescem no máximo até n vértices.",
     False, False),
])

# ════════════════════════════════════════════════════════════════════════════
# 3 METODOLOGIA EXPERIMENTAL
# ════════════════════════════════════════════════════════════════════════════

h1("3", "Metodologia Experimental")

h2("3.1", "Implementações")

body_mixed([
    ("O algoritmo foi implementado em duas linguagens: ", False, False),
    ("Python", True, False),
    (" (python/astar.py) e ", False, False),
    ("JavaScript", True, False),
    (" no Node.js (javascript/astar.js). Ambas usam heap binário com o mesmo critério de "
     "desempate e os mesmos mapeamentos para g-scores, pais e expandidos.", False, False),
])
body_mixed([
    ("Três mecanismos garantem equivalência: (1) o PRNG ", False, False),
    ("mulberry32", True, False),
    (", reproduzido bit a bit nas duas linguagens (aritmética módulo 2³²), garante grades "
     "idênticas; (2) o mesmo desempate garante expansão dos mesmos vértices na mesma ordem; "
     "(3) um script de verificação (tools/verificar_equivalencia.py) confirma via hash FNV-1a. "
     "Resultado: nós expandidos e custos são ", False, False),
    ("idênticos", True, False),
    (" nas duas linguagens em todas as 540 medições.", False, False),
])

h2("3.2", "Cenários de teste e geração das entradas")

body_text(
    "Foram definidos três cenários representando os limites assintóticos do algoritmo:"
)
list_item(
    "Melhor caso — grade vazia: sem obstáculos; origem na célula central da primeira coluna, "
    "destino na central da última coluna (mesma linha). O A* expande exatamente L + 1 ≈ √n + 1 "
    "vértices do caminho reto, sem explorar nenhum outro."
)
list_item(
    "Caso médio — obstáculos aleatórios (25%): células bloqueadas com probabilidade 25% usando "
    "mulberry32 com semente = número da rodada. Origem no canto superior esquerdo, destino no "
    "inferior direito. Se a grade não for conexa, tenta-se a próxima semente."
)
list_item(
    "Pior caso — labirinto em serpentina: paredes horizontais completas nas linhas ímpares, "
    "cada uma com abertura alternando entre extremo direito e esquerdo. A heurística aponta "
    "“para baixo” enquanto o caminho percorre todos os corredores em zigue-zague. "
    "Resultado: ~n/2 vértices expandidos."
)

h2("3.3", "Tamanhos avaliados")

body_mixed([
    ("Cada cenário foi executado em seis tamanhos de lado: ", False, False),
    ("51, 101, 151, 201, 251 e 301", True, False),
    (" células (n = 2.601 a 90.601 células). Os tamanhos nomeados são ", False, False),
    ("pequena (51×51), média (151×151) e grande (301×301)", True, False),
    ("; os três intermediários aumentam a resolução das curvas.", False, False),
])

h2("3.4", "Protocolo de medição")

body_mixed([
    ("Cada combinação de cenário e tamanho foi executada ", False, False),
    ("30 vezes", True, False),
    (" (540 medições por linguagem). A geração das grades fica fora da cronometragem.",
     False, False),
])
body_mixed([
    ("Antes das 30 rodadas, realizaram-se ", False, False),
    ("3 rodadas de aquecimento descartadas", True, False),
    (" em ambas as linguagens — essencial no Node.js, onde o JIT do V8 otimiza o código "
     "nas primeiras execuções.", False, False),
])
body_text(
    "No caso médio cada rodada usa uma grade diferente (semente = número da rodada), de "
    "modo que o desvio-padrão reflete tanto ruído de medição quanto variabilidade real "
    "das instâncias."
)
body_text(
    "Funções de alta precisão: Python: time.perf_counter_ns() (nanossegunos); "
    "JavaScript: process.hrtime.bigint(). Ambas cronometram exclusivamente a chamada do "
    "algoritmo."
)

h2("3.5", "Ambiente de execução")

body_text("Especificações da máquina utilizada:")
list_item("Processador: Intel Core i5-10500 @ 3,10 GHz (6 núcleos / 12 threads).")
list_item("Memória: 16 GB de RAM.")
list_item("Sistema operacional: Windows 11 Pro (build 26200).")
list_item("Interpretador Python: Python 3.13.7 (CPython).")
list_item("Ambiente JavaScript: Node.js v22.21.1 (motor V8).")
body_text(
    "Os benchmarks foram executados de forma isolada: um cenário por vez, em processo "
    "dedicado, sem outras aplicações pesadas em execução concorrente."
)

h2("3.6", "Resumo dos cenários executados")

body_text(
    "As Tabelas 1, 2 e 3 consolidam as estatísticas obtidas (média ± desvio-padrão das "
    "30 rodadas, mínimo e máximo) para os três cenários. Dados completos disponíveis em "
    "results/resultados_python.csv e results/resultados_javascript.csv."
)

caption_table("Tabela 1 – Tempos de execução – Melhor caso (grade vazia, caminho reto) (ms)")
add_table(
    headers=["Linguagem", "Tamanho", "n (células)", "Nós exp.", "Média ± DP (ms)", "Mín (ms)", "Máx (ms)"],
    rows=[
        ["Python",     "pequena", "2.601",  "51",  "0,166 ± 0,004", "0,160", "0,177"],
        ["JavaScript", "pequena", "2.601",  "51",  "0,261 ± 0,205", "0,147", "1,215"],
        ["Python",     "média",   "22.801", "151", "0,704 ± 0,212", "0,522", "1,089"],
        ["JavaScript", "média",   "22.801", "151", "0,379 ± 0,115", "0,300", "0,716"],
        ["Python",     "grande",  "90.601", "301", "1,149 ± 0,051", "1,125", "1,413"],
        ["JavaScript", "grande",  "90.601", "301", "0,338 ± 0,067", "0,259", "0,507"],
    ],
)
source_line()

caption_table("Tabela 2 – Tempos de execução – Caso médio (obstáculos aleatórios 25%) (ms)")
add_table(
    headers=["Linguagem", "Tamanho", "n (células)", "Nós exp.", "Média ± DP (ms)", "Mín (ms)", "Máx (ms)"],
    rows=[
        ["Python",     "pequena", "2.601",  "~251",   "0,669 ± 0,203",  "0,374", "1,107"],
        ["JavaScript", "pequena", "2.601",  "~251",   "0,488 ± 0,168",  "0,244", "0,803"],
        ["Python",     "média",   "22.801", "~1.768", "5,383 ± 5,982",  "1,712", "26,910"],
        ["JavaScript", "média",   "22.801", "~1.768", "1,499 ± 1,943",  "0,392", "8,606"],
        ["Python",     "grande",  "90.601", "~6.756", "22,101 ± 19,169","6,017", "91,736"],
        ["JavaScript", "grande",  "90.601", "~6.756", "9,693 ± 6,448",  "3,021", "31,104"],
    ],
)
source_line()

caption_table("Tabela 3 – Tempos de execução – Pior caso (labirinto em serpentina) (ms)")
add_table(
    headers=["Linguagem", "Tamanho", "n (células)", "Nós exp.", "Média ± DP (ms)", "Mín (ms)", "Máx (ms)"],
    rows=[
        ["Python",     "pequena", "2.601",  "1.301",  "2,235 ± 0,181", "2,115",  "2,986"],
        ["JavaScript", "pequena", "2.601",  "1.301",  "0,351 ± 0,105", "0,297",  "0,791"],
        ["Python",     "média",   "22.801", "11.401", "21,556 ± 0,618","20,700", "23,111"],
        ["JavaScript", "média",   "22.801", "11.401", "4,292 ± 0,543", "3,464",  "5,827"],
        ["Python",     "grande",  "90.601", "45.601", "94,109 ± 2,263","90,638", "98,685"],
        ["JavaScript", "grande",  "90.601", "45.601", "20,874 ± 3,672","16,637", "33,033"],
    ],
)
source_line()

# ════════════════════════════════════════════════════════════════════════════
# 4 RESULTADOS E ANÁLISE DOS GRÁFICOS
# ════════════════════════════════════════════════════════════════════════════

h1("4", "Resultados e Análise dos Gráficos")

body_text(
    "Esta seção apresenta e interpreta todos os gráficos gerados pelos experimentos "
    "(results/graficos/), organizados em quatro grupos: curvas de aderência ao modelo "
    "teórico, análise de casos por linguagem, comparação direta Python × JavaScript e "
    "trabalho algorítmico pelos nós expandidos."
)

h2("4.1", "Curvas de tempo por cenário")

figure("aderencia_pior.png",
       "Figura 1 – Aderência ao modelo teórico: pior caso (labirinto em serpentina)")

body_text(
    "O gráfico em escala log-log apresenta os tempos medidos de Python e JavaScript no pior "
    "caso sobrepostos à curva c·n·log₂(n). Os pontos de ambas as linguagens acompanham essa "
    "reta ao longo de quase duas ordens de grandeza (n = 2.601 a 90.601), confirmando Θ(n log n). "
    "O desvio-padrão é pequeno (Python: ±0,181 ms a ±2,263 ms; JS: ±0,105 ms a ±3,672 ms), "
    "evidenciando medições estáveis. JavaScript é 4,5× mais rápido que Python para n = 90.601."
)

figure("aderencia_medio.png",
       "Figura 2 – Aderência ao modelo teórico: caso médio (obstáculos aleatórios 25%)")

body_text(
    "Os tempos medidos são sobrepostos à curva c·n (crescimento linear). As barras de erro são "
    "consideravelmente maiores do que no pior caso: para n = 90.601, o desvio do Python chega a "
    "19,169 ms sobre média de 22,101 ms; o do JavaScript a 6,448 ms sobre 9,693 ms. Esse "
    "comportamento é esperado: cada rodada usa uma grade sorteada diferente, e o número de nós "
    "expandidos varia entre instâncias. A fração constante de ~8% de células expandidas "
    "confirma o crescimento ≈ linear."
)

figure("aderencia_melhor.png",
       "Figura 3 – Aderência ao modelo teórico: melhor caso (grade vazia)")

body_text(
    "Os pontos são sobrepostos à curva c·√n. O Python segue bem a reta de inclinação 0,5 em "
    "escala log-log, confirmando que o número de nós expandidos (exatamente √n) domina o tempo. "
    "O JavaScript apresenta tempos praticamente constantes (~0,3–0,4 ms), pois o custo fixo do "
    "motor V8 domina quando apenas 51 a 301 nós são expandidos. Para n = 2.601, o Python "
    "(0,166 ms) é mais rápido que o JavaScript (0,261 ms); a situação se inverte a partir "
    "de n = 22.801."
)

figure("casos_python.png",
       "Figura 4 – Análise de casos em Python")

body_text(
    "As três curvas (melhor, médio e pior caso) exibem inclinações distintas em escala log-log: "
    "≈ 0,5 (√n), ≈ 1 (linear) e > 1 (n log n). Para n = 90.601, o pior caso é 82× mais lento "
    "que o melhor caso (94,1 ms contra 1,15 ms). A variabilidade do caso médio é visível pelas "
    "barras de erro, enquanto os cenários determinísticos têm barras muito menores."
)

figure("casos_javascript.png",
       "Figura 5 – Análise de casos em JavaScript")

body_text(
    "O JavaScript exibe a mesma estrutura de três curvas, mas a curva do melhor caso é "
    "praticamente plana. A razão pior/melhor para n = 90.601 é ~62× (20,9 ms contra 0,34 ms), "
    "ligeiramente menor do que no Python (82×), pois o denominador JavaScript é menor nas "
    "entradas grandes. As barras de erro do caso médio também são proporcionalmente grandes."
)

h2("4.2", "Comparação entre Python e JavaScript")

figure("comparacao_linguagens.png",
       "Figura 6 – Comparação direta entre Python e JavaScript (pior caso)")

body_text(
    "O gráfico de barras exibe os tempos médios lado a lado para os três tamanhos nomeados. "
    "A razão Python/JavaScript vai de 6,4× na grade pequena (2,235 ms contra 0,351 ms) a "
    "4,5× na grade grande (94,109 ms contra 20,874 ms). A diferença reflete a natureza dos "
    "runtimes: CPython interpreta bytecodes com despacho dinâmico, enquanto o Node.js (V8) "
    "compila para código de máquina com otimizações JIT agressivas. Apesar disso, "
    "ambas as implementações exibem o mesmo crescimento assintótico, confirmando que "
    "a complexidade é propriedade do algoritmo, não do ambiente de execução."
)

h2("4.3", "Trabalho algorítmico: nós expandidos")

figure("nos_expandidos.png",
       "Figura 7 – Nós expandidos por cenário")

body_text(
    "Este gráfico mede o trabalho algorítmico puro — nós expandidos — em vez de tempo de CPU. "
    "Os valores são idênticos em Python e JavaScript (verificado pelo script de equivalência)."
)
list_item(
    "Melhor caso: proporcional a √n (51, 151, 301 nós). Heurística perfeita; nenhum vértice "
    "além do caminho é expandido."
)
list_item(
    "Caso médio: aproximadamente linear (~251, ~1.768, ~6.756 nós, representando ~8–9% de n). "
    "Fração constante expandida."
)
list_item(
    "Pior caso: proporcional a n/2 (1.301, 11.401, 45.601 nós, ~50% de n). O algoritmo "
    "percorre todos os vértices do único caminho possível."
)

h2("4.4", "Validação do modelo teórico")

body_text(
    "A comparação entre os tempos medidos e as curvas teóricas normalizadas (linhas tracejadas "
    "nas Figuras 1, 2 e 3) demonstra boa aderência empírica:"
)
list_item(
    "Pior caso: curva n log n confirmada ao longo de dois fatores de 10 em n; constante "
    "multiplicativa estável para todos os tamanhos."
)
list_item(
    "Caso médio: reta de inclinação 1 em log-log; maior variância devida à heterogeneidade "
    "das instâncias. Fração constante ~8% de nós sustenta Θ(n)."
)
list_item(
    "Melhor caso: Python segue √n claramente; JavaScript mostra overhead fixo dominante "
    "para n pequeno, convergindo para √n só nas entradas maiores."
)

# ════════════════════════════════════════════════════════════════════════════
# 5 APLICABILIDADE: EFICIÊNCIA E LIMITAÇÕES
# ════════════════════════════════════════════════════════════════════════════

h1("5", "Aplicabilidade: Eficiência e Limitações")

h2("5.1", "Contextos de eficiência e aplicações")

body_text(
    "O A* é o algoritmo de pathfinding mais amplamente utilizado quando: (a) há um único par "
    "origem-destino por consulta; (b) os pesos são não negativos; (c) existe heurística "
    "admissível barata. Nessas condições, o A* supera o Dijkstra: o melhor caso expande "
    "apenas √n vértices, enquanto o Dijkstra exploraria toda a grade (n vértices)."
)
body_text("Principais contextos de aplicação:")
list_item(
    "Jogos e motores de jogo: pathfinding em grades e navmesh (superfícies 3D). Padrão "
    "nas bibliotecas Recast/Detour (Unity, Unreal Engine)."
)
list_item(
    "Robótica e veículos autônomos: planejamento de rota em mapas de ocupação 2D/3D."
)
list_item(
    "Navegação e GPS: algoritmo ALT (A* with Landmarks and Triangle inequality) e "
    "hierarquias de contração para consultas de sub-milissegundo."
)
list_item(
    "Planejamento automático (IA clássica): puzzles como 15-puzzle e cubo mágico com "
    "heurísticas de padrão (pattern databases)."
)
list_item("Redes e telecomunicações: roteamento com estimativas de latência como heurística.")

h2("5.2", "Limitações")

body_text("O A* possui limitações relevantes:")
list_item(
    "Memória O(n): mantém todos os vértices descobertos em memória. Em espaços exponenciais "
    "(≈ 10¹³ estados no 15-puzzle), a fronteira pode esgotar a RAM. Variantes IDA*, SMA* "
    "e fringe search trocam memória por recomputação."
)
list_item(
    "Qualidade da heurística: com h ≡ 0 o algoritmo degenera para Dijkstra, explorando "
    "n vértices sem benefício heurístico — como no pior caso deste trabalho. "
    "Weighted A* explora deliberadamente o trade-off admissibilidade × velocidade."
)
list_item(
    "Ambientes dinâmicos: o A* precisa ser reiniciado a cada mudança do mapa. D* e D* Lite "
    "foram desenvolvidos para replanejamento incremental eficiente."
)
list_item(
    "Pesos negativos: assim como o Dijkstra, o A* não funciona com arestas negativas. "
    "O Bellman-Ford (O(V·E)) é a alternativa."
)
list_item(
    "Múltiplas consultas: contraction hierarchies ou TRANSIT superam o A* puro quando "
    "muitos pares origem-destino precisam ser processados no mesmo grafo."
)
body_text(
    "Em resumo, o A* é altamente eficiente para consulta única em grafo estático com boa "
    "heurística, mas há situações — espaços imensos, ambientes dinâmicos ou múltiplas "
    "consultas — em que outras abordagens são preferíveis."
)

# ════════════════════════════════════════════════════════════════════════════
# 6 REFLEXÃO TEÓRICA: CLASSES DE COMPLEXIDADE
# ════════════════════════════════════════════════════════════════════════════

h1("6", "Reflexão Teórica: Classes de Complexidade")

body_text(
    "Esta seção relaciona o algoritmo estudado com as classes de complexidade P e NP e "
    "discute problemas correlatos de difícil solução."
)

h2("6.1", "O problema pertence à classe P?")

body_text(
    "Sim. A classe P reúne os problemas de decisão resolúveis por algoritmo determinístico em "
    "tempo polinomial. O problema do caminho mínimo — versão de decisão: “existe caminho de "
    "custo ≤ k?” — é resolvido pelo A* em O(n log n), que é polinomial em n. Portanto o "
    "problema pertence a P."
)
body_text(
    "Como P ⊆ NP, o problema também pertence a NP: dado qualquer caminho proposto, basta "
    "percorrê-lo uma vez (O(n)) para verificar se o custo é ≤ k."
)

h2("6.2", "Existe uma versão NP do problema?")

body_text(
    "O caminho mínimo usual não é NP-difícil, mas variações aparentemente pequenas elevam "
    "drasticamente a dificuldade:"
)
list_item(
    "Caminho simples mais longo (longest path): NP-difícil. Trocar “menor” por “maior” "
    "elimina a subestrutura ótima que torna o caminho mínimo polinomial."
)
list_item(
    "Caminho mínimo com restrições de recursos (RCSP): NP-difícil mesmo com pesos positivos. "
    "A restrição de orçamento adicional quebra a estrutura polinomial."
)
list_item(
    "Planejamento multiagente (MAPF): NP-difícil. Encontrar caminhos de custo total mínimo "
    "para múltiplos agentes sem colisões não admite solução polinomial geral."
)
list_item(
    "Problema do (n²−1)-puzzle ótimo: NP-difícil (Ratner e Warmuth, 1990). O A* é ótimo "
    "nesses espaços, mas o número de estados é exponencial — logo o tempo também."
)
body_text(
    "A fronteira entre tratável e intratável é tênue: caminho mínimo em grafos explícitos "
    "está em P, mas restrições adicionais, troca do critério de otimização ou "
    "exponencialidade do espaço de estados mudam a classe."
)

h2("6.3", "Problemas semelhantes NP-completos")

body_text(
    "Vários problemas de grafos se assemelham ao caminho mínimo mas são NP-completos:"
)
list_item(
    "Problema do Caixeiro-Viajante (TSP): encontrar o ciclo de menor custo que visita todos "
    "os vértices exatamente uma vez. A versão de decisão (“existe tal ciclo com custo ≤ k?”) "
    "é NP-completa. A diferença fundamental: o TSP exige caminho hamiltoniano, o que quebra "
    "a propriedade de subestrutura ótima."
)
list_item(
    "Caminho Hamiltoniano: determinar se existe caminho simples passando por todos os vértices "
    "exatamente uma vez é NP-completo. O TSP com grafo completo é generalização ponderada."
)
list_item(
    "Múltiplos agentes com orçamento compartilhado: variantes em que n agentes minimizam "
    "custo total com restrição de orçamento são NP-difíceis."
)
body_text(
    "A comparação é instrutiva: o A* em P escala até 90.601 células em < 100 ms, enquanto "
    "o TSP e o caminho hamiltoniano não têm solução polinomial conhecida. O fato de o A* "
    "pertencer a P é uma vantagem teórica expressiva: eficiência garantida e escalabilidade "
    "comprovada, ao contrário dos problemas NP-completos correlatos. Caso fosse descoberto "
    "algoritmo polinomial para qualquer NP-completo, ter-se-ia P = NP — uma das maiores "
    "questões em aberto da Ciência da Computação."
)

# ════════════════════════════════════════════════════════════════════════════
# 7 EXECUÇÃO DO PROJETO
# ════════════════════════════════════════════════════════════════════════════

h1("7", "Execução do Projeto")

body_mixed([
    ("Esta seção descreve como reproduzir o experimento a partir do código-fonte. "
     "Repositório: ", False, False),
    ("https://github.com/P4d1lh4/Projeto-Teoria-A-", False, True),
])
body_text("Estrutura de pastas:")
code_block([
    "python/         código + benchmark em Python",
    "javascript/     código + benchmark em JavaScript (Node.js)",
    "tools/          script de verificação de equivalência",
    "results/        CSVs e gráficos gerados pelos experimentos",
    "docs/           relatório, slides e tabelas",
    "demo/           demonstração visual interativa (index.html)",
])

h2("7.1", "Pré-requisitos")

body_text(
    "Python 3.10+ e Node.js 18+. Única dependência Python de análise:"
)
code_block(["pip install matplotlib"])

h2("7.2", "Executando os testes")

code_block([
    "# Testes unitários Python",
    "python -m unittest discover -s python -v",
    "",
    "# Testes JavaScript",
    "node --test javascript/test/astar.test.js",
    "",
    "# Verificação de equivalência entre linguagens",
    "python tools/verificar_equivalencia.py",
])

h2("7.3", "Executando os benchmarks")

code_block([
    "# Benchmark Python  →  results/resultados_python.csv",
    "python python/benchmark.py",
    "",
    "# Benchmark JavaScript  →  results/resultados_javascript.csv",
    "node javascript/benchmark.js",
    "",
    "# Modo rápido (5 rodadas, tamanhos nomeados)",
    "python python/benchmark.py --rapido",
])

h2("7.4", "Gerando os gráficos e tabelas")

body_text("Após ambos os benchmarks:")
code_block(["python python/charts.py"])
body_text(
    "Saídas: results/graficos/ (PNGs) e docs/tabelas_resultados.md (tabelas Markdown)."
)

h2("7.5", "Demonstração visual")

body_text(
    "Abrir demo/index.html em qualquer navegador exibe o A* explorando a grade em tempo "
    "real, com visualização das células abertas, fechadas e do caminho encontrado — útil "
    "para a apresentação oral."
)

# ════════════════════════════════════════════════════════════════════════════
# 8 CONCLUSÃO
# ════════════════════════════════════════════════════════════════════════════

h1("8", "Conclusão")

body_text(
    "Este trabalho apresentou um estudo teórico e experimental do algoritmo A* para o problema "
    "do caminho mínimo em grades 2D com obstáculos. Teoricamente, mostrou-se que o algoritmo, "
    "implementado com heap binário de mínimo e heurística de Manhattan admissível e consistente, "
    "é ótimo, tem limite superior O(n log n), limite inferior Ω(√n·log n) no melhor caso e não "
    "admite Θ global. Por ser polinomial, pertence à classe P — ao contrário de variantes "
    "correlatas como TSP e caminho hamiltoniano, que são NP-completos."
)
body_text(
    "Os experimentos confirmaram as previsões: no pior caso os tempos crescem como n log n; "
    "no caso médio o crescimento é ≈ linear com ~8% das células expandidas; no melhor caso "
    "exatamente √n vértices são expandidos. Python segue √n claramente; JavaScript exibe "
    "overhead fixo dominante para n pequeno."
)
body_text(
    "Na comparação entre linguagens, JavaScript (Node.js/V8) foi de 4,5× a 6,4× mais rápido "
    "no pior caso graças ao JIT. Python apresentou menor variância nos cenários determinísticos "
    "e foi mais eficiente para entradas muito pequenas no melhor caso. Ambas as implementações "
    "exibiram o mesmo padrão assintótico, confirmando que complexidade é propriedade do "
    "algoritmo, não da linguagem."
)
body_text(
    "A verificação formal de equivalência — mesmo PRNG bit a bit, mesmo critério de desempate "
    "e script FNV-1a — garantiu que os dois runtimes processaram exatamente as mesmas entradas "
    "e expandiram exatamente os mesmos nós, tornando a comparação de tempos uma medição pura "
    "de desempenho de runtime, sem interferência algorítmica."
)

# ════════════════════════════════════════════════════════════════════════════
# REFERÊNCIAS
# ════════════════════════════════════════════════════════════════════════════

h1("", "Referências")   # sem número — ABNT

def ref(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(p, line=LS_10, before=0, after=160)
    _indent(p, hanging=FIRST)
    _run(p, text)
    return p

ref(
    "HART, P. E.; NILSSON, N. J.; RAPHAEL, B. A formal basis for the heuristic determination "
    "of minimum cost paths. IEEE Transactions on Systems Science and Cybernetics, v. 4, n. 2, "
    "p. 100–107, 1968."
)
ref(
    "RUSSELL, S.; NORVIG, P. Artificial Intelligence: A Modern Approach. 4. ed. Hoboken: "
    "Pearson, 2020. (Capítulo 3 — busca informada e exploração)"
)
ref(
    "CORMEN, T. H. et al. Introduction to Algorithms. 4. ed. Cambridge: MIT Press, 2022. "
    "(Capítulo 22 — Dijkstra e filas de prioridade)"
)
ref(
    "RATNER, D.; WARMUTH, M. The (n²−1)-puzzle and related relocation problems. Journal of "
    "Symbolic Computation, v. 10, n. 2, p. 111–137, 1990."
)
ref(
    "GAREY, M. R.; JOHNSON, D. S. Computers and Intractability: A Guide to the Theory of "
    "NP-Completeness. New York: W. H. Freeman, 1979."
)
ref(
    "ZIVIANI, N. Projeto de Algoritmos com Implementações em Java e C++. São Paulo: "
    "Cengage Learning, 2007."
)

# ════════════════════════════════════════════════════════════════════════════
doc.save(str(SAIDA))
print(f"✓  {SAIDA}")
